"""合同解析服务"""
import io
import json
import re
import logging
import time
import base64
import asyncio
import requests as http_requests
from openai import OpenAI
from src.core.config import settings
from src.contract_parser.prompts import SYSTEM_PROMPT, build_user_prompt

logger = logging.getLogger(__name__)

# document-converter 服务地址（K8s 内部 DNS）
CONVERTER_URL = "http://document-converter.cronmail.svc.cluster.local:8080"


def convert_to_pdf(file_bytes: bytes, filename: str) -> bytes:
    """通过 document-converter 服务将文件转为 PDF"""
    resp = http_requests.post(
        f"{CONVERTER_URL}/convert",
        files={'file': (filename, io.BytesIO(file_bytes))},
        timeout=120,
    )
    if resp.status_code != 200:
        raise RuntimeError(f"文档转换失败: HTTP {resp.status_code}")
    return resp.content


def extract_text_from_docx(file_bytes: bytes) -> str:
    """从 Word 文件提取纯文本（含表格）"""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    
    # 1. 段落文本
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    
    # 2. 表格数据（服务清单等关键信息在这里）
    for ti, table in enumerate(doc.tables):
        parts.append(f"\n[表格{ti+1}]")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            parts.append(" | ".join(cells))
    
    return "\n".join(parts)


def extract_images_from_pdf(file_bytes: bytes, max_pages: int = 50) -> list[str]:
    """将 PDF 逐页转为 base64 图片列表（JPEG 压缩 + 并行）"""
    from pdf2image import convert_from_bytes
    from concurrent.futures import ThreadPoolExecutor
    
    images = convert_from_bytes(file_bytes, dpi=100, fmt='jpeg', thread_count=4)
    if len(images) > max_pages:
        images = images[:max_pages]
    
    def _encode(img) -> str:
        buf = io.BytesIO()
        img.save(buf, format='JPEG', quality=85, optimize=True)
        buf.seek(0)
        return base64.b64encode(buf.read()).decode('utf-8')
    
    with ThreadPoolExecutor(max_workers=4) as pool:
        result = list(pool.map(_encode, images))
    
    return result


def extract_text_from_doc(file_bytes: bytes) -> str:
    """从旧版 .doc 文件提取纯文本（通过 antiword）"""
    import tempfile
    import subprocess
    with tempfile.NamedTemporaryFile(suffix='.doc', delete=True) as tmp:
        tmp.write(file_bytes)
        tmp.flush()
        try:
            result = subprocess.run(
                ['antiword', tmp.name],
                capture_output=True, text=True, timeout=30,
            )
            text = result.stdout.strip()
            if not text:
                raise ValueError("antiword 无法提取文本，请将文件另存为 .docx 或 .pdf 格式")
            return text
        except FileNotFoundError:
            raise ValueError("antiword 未安装，请联系管理员")

def parse_contract(file_bytes: bytes, filename: str, contract_type: str) -> dict:
    """解析合同文件 — 所有格式统一走 Vision 多模态图片识别

    Raises:
        ValueError: 文件格式不支持
        RuntimeError: 文档转换或 LLM 调用失败
    """
    processing_info = {"file_size_kb": round(len(file_bytes) / 1024, 1)}

    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

    # .doc/.docx → converter → PDF → 统一走 Vision
    if ext in ('doc', 'docx'):
        processing_info["mode"] = "vision"
        processing_info["file_type"] = ext
        try:
            pdf_bytes = convert_to_pdf(file_bytes, filename)
        except Exception as e:
            raise RuntimeError(f"文档转换失败: {e}")
        processing_info["converted_from"] = ext
        return parse_contract_vision(pdf_bytes, filename.rsplit('.', 1)[0] + '.pdf', contract_type, processing_info)

    # .pdf → 直接走 Vision
    elif ext == 'pdf':
        processing_info["mode"] = "vision"
        processing_info["file_type"] = "pdf"
        return parse_contract_vision(file_bytes, filename, contract_type, processing_info)

    # .png/.jpg/.jpeg → 包装为单页 PDF 走 Vision
    elif ext in ('png', 'jpg', 'jpeg'):
        from PIL import Image
        import io
        img = Image.open(io.BytesIO(file_bytes))
        if img.mode in ('RGBA', 'P'):
            img = img.convert('RGB')
        pdf_buffer = io.BytesIO()
        img.save(pdf_buffer, format='PDF')
        pdf_bytes = pdf_buffer.getvalue()
        processing_info["mode"] = "vision"
        processing_info["file_type"] = ext
        processing_info["converted_from"] = ext
        return parse_contract_vision(pdf_bytes, filename.rsplit('.', 1)[0] + '.pdf', contract_type, processing_info)

    else:
        raise ValueError(f"不支持的文件格式: .{ext}，支持 .doc / .docx / .pdf / .png / .jpg / .jpeg")


def parse_contract_vision(file_bytes: bytes, filename: str, contract_type: str, processing_info: dict = None) -> dict:
    """两阶段管道：
    阶段 1：分批 OCR 提取纯文字 → 拼接为完整 MD
    阶段 2：文本 LLM 从全文提取结构化 JSON
    """
    t0 = time.time()
    timing = {}
    if processing_info is None:
        processing_info = {"mode": "vision", "file_size_kb": round(len(file_bytes) / 1024, 1), "file_type": "pdf"}

    # 1. PDF 转图片
    t1 = time.time()
    b64_images = extract_images_from_pdf(file_bytes)
    if not b64_images:
        raise ValueError("PDF 无有效页面")
    timing["pdf_to_images"] = {"seconds": round(time.time() - t1, 1), "pages": len(b64_images)}
    processing_info["pdf_pages"] = len(b64_images)
    processing_info["extract_seconds"] = timing["pdf_to_images"]["seconds"]

    # ============================================================
    # 阶段 1：分批 OCR 提取纯文字
    # ============================================================
    from src.contract_parser.prompts import OCR_SYSTEM_PROMPT, get_ocr_batch_prompt

    client = OpenAI(
        base_url=settings.LLM_BASE_URL,
        api_key=settings.LLM_API_KEY,
        timeout=300,
    )

    BATCH_SIZE = 3
    total_pages = len(b64_images)
    batches = _chunk_list(b64_images, BATCH_SIZE)
    total_batches = len(batches)
    ocr_parts = []  # 每批 OCR 结果
    ocr_timing = []

    for batch_idx, batch_images in enumerate(batches):
        t_batch = time.time()
        batch_num = batch_idx + 1
        start_page = batch_idx * BATCH_SIZE + 1
        end_page = min(start_page + len(batch_images) - 1, total_pages)

        batch_prompt = get_ocr_batch_prompt(batch_num, total_batches, start_page, end_page)
        content = [{"type": "text", "text": batch_prompt}]
        for b64 in batch_images:
            content.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}})

        try:
            response = client.chat.completions.create(
                model=settings.LLM_MODEL,
                messages=[
                    {"role": "system", "content": OCR_SYSTEM_PROMPT},
                    {"role": "user", "content": content},
                ],
                temperature=0.3,
                max_tokens=16384,
                extra_body={"chat_template_kwargs": {"enable_thinking": True}},
            )
            raw = response.choices[0].message.content or ""
            ocr_parts.append(raw)
            logger.info(f"OCR batch {batch_num}/{total_batches} (pages {start_page}-{end_page}) len={len(raw)}")
        except Exception as e:
            logger.error(f"OCR batch {batch_num} failed: {e}")
            ocr_parts.append(f"--- 第{start_page}-{end_page}页（OCR 失败） ---")

        elapsed = round(time.time() - t_batch, 1)
        ocr_timing.append({"batch": batch_num, "pages": f"{start_page}-{end_page}", "seconds": elapsed})

    # 拼接完整文字
    full_text = "\n\n".join(ocr_parts)
    logger.info(f"OCR complete: {len(full_text)} chars from {total_pages} pages")

    # ============================================================
    # 阶段 2：文本 LLM 从全文提取结构化 JSON
    # ============================================================
    from src.contract_parser.prompts import EXTRACT_SYSTEM_PROMPT, get_extract_prompt

    t_extract = time.time()
    extract_prompt = get_extract_prompt(contract_type, full_text)

    try:
        response = client.chat.completions.create(
            model=settings.LLM_MODEL,
            messages=[
                {"role": "system", "content": EXTRACT_SYSTEM_PROMPT},
                {"role": "user", "content": extract_prompt},
            ],
            temperature=0.3,
            max_tokens=32768,
            extra_body={"chat_template_kwargs": {"enable_thinking": True}},
        )
        raw = response.choices[0].message.content
        result = _parse_llm_json(raw)
        import sys
        print(f"[contract_parser] [sync] Extract keys: {list(result.keys()) if result else 'EMPTY'}, party_a={result.get('party_a_name')}, party_b={result.get('party_b_name')}", file=sys.stderr)
        rs = result.get("resource_summary", {})
        logger.info(f"[sync] Extract result: resource_summary items={rs.get('items')}, stats={rs.get('stats')}")
    except Exception as e:
        logger.error(f"Extract LLM failed: {e}")
        result = {}

    timing["extract"] = {"seconds": round(time.time() - t_extract, 1)}

    # 后处理：代码层过滤确认单 + 去重 + 资源统计
    _post_process_tables(result)

    timing["total_vision"] = {"seconds": round(time.time() - t0, 1)}
    result["_processing_info"] = processing_info
    result["_timing"] = timing
    result["_ocr_text"] = full_text  # 调试用，后续可移除
    return result


def _chunk_list(lst: list, chunk_size: int) -> list:
    """将列表按 chunk_size 分批"""
    return [lst[i:i + chunk_size] for i in range(0, len(lst), chunk_size)]


def parse_payment_receipt(file_bytes: bytes, filename: str) -> dict:
    """解析回款凭证（回执单/发票）——简化版 Vision 管道，不开启思考

    流程：
    1. 文件转 PDF（如需）
    2. PDF 逐页拆图
    3. Vision LLM 逐批 OCR（enable_thinking=False）
    4. 拼接文本 → 文本 LLM 提取金额/日期

    Returns:
        {"amount": "...", "payment_date": "...", "payer": "...", "doc_type": "..."}
    """
    from src.contract_parser.prompts import (
        PAYMENT_RECEIPT_SYSTEM_PROMPT, PAYMENT_RECEIPT_EXTRACT_PROMPT,
        get_ocr_batch_prompt,
    )

    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    # 获取 PDF bytes
    if ext == 'pdf':
        pdf_bytes = file_bytes
    elif ext in ('png', 'jpg', 'jpeg'):
        from PIL import Image as PILImage
        img = PILImage.open(io.BytesIO(file_bytes))
        if img.mode in ('RGBA', 'P'):
            img = img.convert('RGB')
        pdf_buf = io.BytesIO()
        img.save(pdf_buf, format='PDF')
        pdf_bytes = pdf_buf.getvalue()
    else:
        pdf_bytes = convert_to_pdf(file_bytes, filename)

    # 拆页：使用 extract_images_from_pdf（接受 file_bytes，返回 base64 字符串列表）
    images_b64 = extract_images_from_pdf(pdf_bytes, max_pages=10)
    total_pages = len(images_b64)
    if total_pages == 0:
        return {"amount": None, "payment_date": None, "payer": None, "doc_type": "unknown"}

    ocr_parts: list[str] = []
    batch_size = 3
    batches = list(_chunk_list(images_b64, batch_size))
    total_batches = len(batches)

    client = OpenAI(api_key=settings.LLM_API_KEY, base_url=settings.LLM_BASE_URL)

    for batch_num, batch in enumerate(batches, 1):
        start_page = (batch_num - 1) * batch_size + 1
        end_page = min(batch_num * batch_size, total_pages)
        content = [{"type": "text", "text": get_ocr_batch_prompt(batch_num, total_batches, start_page, end_page)}]
        for img_b64 in batch:
            content.append({"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"}})

        response = client.chat.completions.create(
            model=settings.LLM_MODEL,
            messages=[{"role": "user", "content": content}],
            temperature=0.1,
            max_tokens=8192,
            extra_body={"chat_template_kwargs": {"enable_thinking": False}},
        )
        ocr_parts.append(response.choices[0].message.content or "")

    full_text = "\n\n".join(ocr_parts)
    if len(full_text) > 32000:
        full_text = full_text[:32000]

    # 阶段 2：提取
    extract_prompt = PAYMENT_RECEIPT_EXTRACT_PROMPT.format(full_text=full_text)
    response = client.chat.completions.create(
        model=settings.LLM_MODEL,
        messages=[
            {"role": "system", "content": PAYMENT_RECEIPT_SYSTEM_PROMPT},
            {"role": "user", "content": extract_prompt},
        ],
        temperature=0.1,
        max_tokens=4096,
        extra_body={"chat_template_kwargs": {"enable_thinking": False}},
    )
    raw = response.choices[0].message.content
    return _parse_llm_json(raw) or {}
def _post_process_tables(accumulated):
    """后处理：过滤非交付表格 + 去重"""
    if isinstance(accumulated, list):
        return  # LLM 返回了数组而非对象，跳过
    if not accumulated.get("raw_tables"):
        return

    # 关键词黑名单（title 包含任一关键词的表格直接丢弃）
    BLACKLIST_KEYWORDS = ["确认单", "验收单", "验收报告", "签字页", "签章页", "审批单", "确认函"]
    
    filtered = []
    for t in accumulated["raw_tables"]:
        if not isinstance(t, dict):
            continue
        title = t.get("title", "") or ""
        if any(kw in title for kw in BLACKLIST_KEYWORDS):
            logger.info(f"后处理过滤表格: {title}")
            continue
        filtered.append(t)
    
    # 去重：按 (title, headers) 保留行数最多的版本
    deduped = {}
    for t in filtered:
        key = (t.get("title", ""), tuple(t.get("headers", [])))
        if key not in deduped or len(t.get("rows", [])) > len(deduped[key].get("rows", [])):
            deduped[key] = t
    accumulated["raw_tables"] = list(deduped.values())

    # 确保 resource_summary 存在
    if "resource_summary" not in accumulated:
        accumulated["resource_summary"] = {"stats": {}, "summary_text": ""}

    # 从 items 精确计算 stats（代码层做乘法累加，不依赖 Agent 算术）
    items = accumulated["resource_summary"].get("items", [])
    if items:
        total = {"vcpu": 0, "memory_gb": 0, "storage_gb": 0, "gpu_count": 0, "gpu_tops": 0}
        for item in items:
            qty = item.get("qty", 0) or 0
            for key in total:
                val = item.get(key, 0) or 0
                total[key] += val * qty
        accumulated["resource_summary"]["stats"] = total
        logger.info(f"资源统计（代码层精确计算）: items={items}, total={total}")
    else:
        logger.warning(f"资源统计: resource_summary 中无 items，stats={accumulated['resource_summary'].get('stats', {})}")


def _extract_number_from_row(row: list) -> int:
    """从表格行中提取数量（最后一列纯数字）"""
    if not row:
        return 0
    for cell in reversed(row):
        cell = str(cell).strip().replace(',', '').replace('个', '').replace('台', '')
        try:
            return int(cell)
        except (ValueError, TypeError):
            continue
    return 0


def _parse_llm_json(raw: str) -> dict:
    """安全解析 LLM 返回的 JSON"""
    if not raw or not raw.strip():
        return {}
    try:
        cleaned = re.sub(r'^```(?:json)?\s*', '', raw.strip())
        cleaned = re.sub(r'\s*```$', '', cleaned)
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return {}


# 字符串字段：首次非空有效值优先，防止后期批次用确认单内容污染
_STRING_FIELDS = {"name", "contract_no", "party_a_name", "party_b_name", "amount",
                  "start_date", "end_date", "project_name", "contract_content",
                  "delivery_requirements", "contract_type", "process_records"}

def _deep_merge(base: dict, update: dict):
    """深度合并。标量字段首次有效值优先，dict 递归合并，list extend。
    
    策略：字符串字段一旦有值就不覆盖，防止后续批次的附件页污染正文信息。
    remark 字段除外，允许追加。
    """
    for k, v in update.items():
        if v is None or v == [] or v == {}:
            continue
        if isinstance(v, dict) and isinstance(base.get(k), dict):
            _deep_merge(base[k], v)
        elif isinstance(v, list) and isinstance(base.get(k), list):
            base[k].extend(v)
        elif k in _STRING_FIELDS:
            # 字符串关键字段：首次有效值优先
            if not base.get(k):
                base[k] = v
        else:
            # remark、resource_summary 等：允许覆盖
            base[k] = v


async def parse_contract_stream(file_bytes: bytes, filename: str, contract_type: str):
    """SSE 流式解析：阶段 1 分批 OCR → 阶段 2 文本汇总提取"""
    t0 = time.time()
    processing_info = {"mode": "vision", "file_size_kb": round(len(file_bytes) / 1024, 1), "file_type": "pdf"}

    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

    if ext in ('doc', 'docx'):
        processing_info["file_type"] = ext
        loop = asyncio.get_running_loop()
        pdf_bytes = await loop.run_in_executor(None, convert_to_pdf, file_bytes, filename)
        processing_info["converted_from"] = ext
    else:
        pdf_bytes = file_bytes

    # 1. PDF 转图片
    t1 = time.time()
    loop = asyncio.get_running_loop()
    b64_images = await loop.run_in_executor(None, extract_images_from_pdf, pdf_bytes)
    total_pages = len(b64_images)
    yield {"event": "progress", "data": json.dumps({
        "step": "pdf_to_images", "pages": total_pages, "seconds": round(time.time() - t1, 1)
    }, ensure_ascii=False)}

    # ============================================================
    # 阶段 1：分批 OCR 提取纯文字
    # ============================================================
    from src.contract_parser.prompts import OCR_SYSTEM_PROMPT, get_ocr_batch_prompt

    client = OpenAI(
        base_url=settings.LLM_BASE_URL,
        api_key=settings.LLM_API_KEY,
        timeout=300,
    )

    BATCH_SIZE = 3
    batches = _chunk_list(b64_images, BATCH_SIZE)
    total_batches = len(batches)
    ocr_parts = []

    for batch_idx, batch_images in enumerate(batches):
        t_batch = time.time()
        batch_num = batch_idx + 1
        start_page = batch_idx * BATCH_SIZE + 1
        end_page = min(start_page + len(batch_images) - 1, total_pages)

        batch_prompt = get_ocr_batch_prompt(batch_num, total_batches, start_page, end_page)
        content = [{"type": "text", "text": batch_prompt}]
        for b64 in batch_images:
            content.append({"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}})

        try:
            response = await loop.run_in_executor(
                None,
                lambda content=content, sp=OCR_SYSTEM_PROMPT, m=settings.LLM_MODEL: client.chat.completions.create(
                    model=m,
                    messages=[{"role": "system", "content": sp}, {"role": "user", "content": content}],
                    temperature=0.1, max_tokens=16384,
                    extra_body={"chat_template_kwargs": {"enable_thinking": False}},
                )
            )
            raw = response.choices[0].message.content or ""
            ocr_parts.append(raw)
            logger.info(f"OCR batch {batch_num}/{total_batches} (pages {start_page}-{end_page}) len={len(raw)}")
        except Exception as e:
            logger.error(f"OCR batch {batch_num} failed: {e}")
            ocr_parts.append(f"--- 第{start_page}-{end_page}页（OCR 失败） ---")

        elapsed = round(time.time() - t_batch, 1)
        pageRange = str(start_page) if start_page == end_page else f"{start_page}-{end_page}"
        yield {"event": "batch", "data": json.dumps({
            "batch": batch_num, "total_batches": total_batches,
            "start_page": start_page, "end_page": end_page,
            "seconds": elapsed, "found_fields": ["OCR 文字提取"],
            "image_base64": batch_images[0],
        }, ensure_ascii=False)}

    # 拼接完整文字
    full_text = "\n\n".join(ocr_parts)
    logger.info(f"OCR complete: {len(full_text)} chars from {total_pages} pages")
    yield {"event": "progress", "data": json.dumps({
        "step": "ocr_done", "chars": len(full_text), "seconds": round(time.time() - t1, 1)
    }, ensure_ascii=False)}

    # ============================================================
    # 阶段 2：文本 LLM 从全文提取结构化 JSON
    # ============================================================
    from src.contract_parser.prompts import EXTRACT_SYSTEM_PROMPT, get_extract_prompt

    t_extract = time.time()
    extract_prompt = get_extract_prompt(contract_type, full_text)

    try:
        response = await loop.run_in_executor(
            None,
            lambda ep=extract_prompt, sp=EXTRACT_SYSTEM_PROMPT, m=settings.LLM_MODEL: client.chat.completions.create(
                model=m,
                messages=[{"role": "system", "content": sp}, {"role": "user", "content": ep}],
                temperature=0.3, max_tokens=32768,
                extra_body={"chat_template_kwargs": {"enable_thinking": True}},
            )
        )
        raw = response.choices[0].message.content
        result = _parse_llm_json(raw)
        import sys
        print(f"[contract_parser] [stream] Extract keys: {list(result.keys()) if result else 'EMPTY'}, party_a={result.get('party_a_name')}, party_b={result.get('party_b_name')}", file=sys.stderr)
        rs = result.get("resource_summary", {})
        logger.info(f"[stream] Extract result: resource_summary items={rs.get('items')}, stats={rs.get('stats')}")
    except Exception as e:
        logger.error(f"Extract LLM failed: {e}")
        result = {}

    yield {"event": "progress", "data": json.dumps({
        "step": "extract_done", "seconds": round(time.time() - t_extract, 1)
    }, ensure_ascii=False)}

    # 后处理：过滤 + 去重 + 资源统计
    _post_process_tables(result)

    result["_processing_info"] = processing_info
    result["_processing_info"]["elapsed_seconds"] = round(time.time() - t0, 1)
    result["_ocr_text"] = full_text
    if result.get("raw_tables"):
        result["raw_tables_json"] = json.dumps(result["raw_tables"], ensure_ascii=False)
    yield {"event": "done", "data": json.dumps({"fields": result}, ensure_ascii=False)}


# ============================================================
# 辅助函数（_deep_merge 已不再需要，保留兼容）
# ============================================================

def _deep_merge(base: dict, update: dict):
    """深度合并（已废弃，两阶段管道不再使用）"""
    for k, v in update.items():
        if v is None or v == [] or v == {}:
            continue
        if isinstance(v, dict) and isinstance(base.get(k), dict):
            _deep_merge(base[k], v)
        elif isinstance(v, list) and isinstance(base.get(k), list):
            base[k].extend(v)
        else:
            base[k] = v
